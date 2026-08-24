"""Extract text from PDF, DOCX, and TXT files."""

import io
import logging
import zipfile
from pathlib import Path
from typing import Iterator

logger = logging.getLogger(__name__)

# Folder-name → document category mapping (matches documents.zip structure)
FOLDER_CATEGORY_MAP = {
    "kawasan_wilayah":   "regional_profile",
    "komoditas_standar": "technical_standards",
    "program_layanan":   "service_programs",
    "regulasi_kebijakan":"regulations_policies",
    "umum_referensi":    "general_reference",
}


class DocumentProcessor:
    """Extract raw text from PDF/DOCX/TXT files inside documents.zip."""

    def __init__(self, zip_path: str | Path):
        self.zip_path = Path(zip_path)

    def iter_documents(self) -> Iterator[dict]:
        """
        Yield dicts with keys:
          filename, category, doc_type, text
        """
        with zipfile.ZipFile(self.zip_path, "r") as zf:
            for entry in zf.infolist():
                if entry.is_dir():
                    continue
                path = Path(entry.filename)
                folder = path.parent.name
                category = FOLDER_CATEGORY_MAP.get(folder, "unknown")
                suffix = path.suffix.lower()

                try:
                    data = zf.read(entry.filename)
                    if suffix == ".txt":
                        text = data.decode("utf-8", errors="replace")
                    elif suffix == ".pdf":
                        text = self._extract_pdf(data)
                    elif suffix == ".docx":
                        text = self._extract_docx(data)
                    else:
                        continue

                    text = text.strip()
                    if len(text) < 50:
                        continue

                    yield {
                        "filename": path.name,
                        "category": category,
                        "doc_type": suffix.lstrip("."),
                        "text": text,
                    }
                except Exception as exc:
                    logger.warning(f"Failed to process {entry.filename}: {exc}")

    def _extract_pdf(self, data: bytes) -> str:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                # Also extract tables as text
                tables = page.extract_tables() or []
                for table in tables:
                    for row in table:
                        row_text = " | ".join(str(c) for c in row if c)
                        if row_text.strip():
                            text_parts.append(row_text)
                if page_text.strip():
                    text_parts.append(page_text)
        return "\n".join(text_parts)

    def _extract_docx(self, data: bytes) -> str:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _extract_metadata(self, text: str, filename: str) -> dict:
        """Heuristically extract domain metadata from document text."""
        import re
        meta = {}

        # Region name: often first line or after "KAWASAN"
        region_match = re.search(r"KAWASAN\s+(?:TRANSMIGRASI\s+)?([A-Z][A-Z\s\-]+?)(?:\s*[–\-]|\s*,|\n)", text)
        if region_match:
            meta["region"] = region_match.group(1).strip().title()

        # Province
        prov_match = re.search(r"(?:Provinsi|PROVINSI|Sulawesi|Kalimantan|Papua|Sumatera|Jawa|Nusa Tenggara|Maluku)\s*([A-Z][a-z]+(?:\s[A-Z][a-z]+)?)", text)
        if prov_match:
            meta["province"] = prov_match.group(0).strip()

        # Year
        year_match = re.search(r"\b(20\d{2})\b", text)
        if year_match:
            meta["year"] = year_match.group(1)

        # Commodity keywords
        commodities = []
        for c in ["padi", "jagung", "kopi", "kakao", "kelapa sawit", "karet", "sapi", "udang"]:
            if c in text.lower():
                commodities.append(c)
        if commodities:
            meta["commodities"] = commodities

        # Regulation type
        if any(k in text.lower() for k in ["peraturan pemerintah", "pp nomor"]):
            meta["regulation_type"] = "PP"
        elif any(k in text.lower() for k in ["peraturan menteri", "permen"]):
            meta["regulation_type"] = "Permen"
        elif any(k in text.lower() for k in ["undang-undang", "uu nomor"]):
            meta["regulation_type"] = "UU"

        return meta
